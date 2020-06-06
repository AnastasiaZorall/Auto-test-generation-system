import random
from docx import Document
from docx.shared import Pt
from tkinter import *
from tkinter.ttk import Checkbutton


def variat(it):
    """
    Function allow to work with file of variant, which you need right now
    :param it: The number of varial
    :return:
    """
    global dti, dti1
    if it == 0:
        dti = dt_1
        dti1 = dt1_1
    if it == 1:
        dti = dt_2
        dti1 = dt1_2
    if it == 2:
        dti = dt_3
        dti1 = dt1_3
    if it == 3:
        dti = dt_4
        dti1 = dt1_4
    if it == 4:
        dti = dt_5
        dti1 = dt1_5


def fio_add(dti):
    """
    Function add first/second name and patronymic in file
    No input and output data
    """
    paragraph = dti.add_paragraph('Фамилия, Имя, Группа' + '_' * 72)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(3.0)
    paragraph = dti.add_paragraph('Дата выполнения' + '_' * 80 + '\n')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2.0)


def input_file_docx(str_write, str_answer):
    """
    Function add a new paragraph in two files (main+answers)
    :param str_write: string which will be add to the main file
    :param str_answer: string which will be add to file with answers
    """
    paragraph = dti.add_paragraph(str_write)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(1.0)

    paragraph = dti1.add_paragraph(str_answer)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(1.0)


def only_int(p):
    """
    Checking string p if it is a number
    :param p: string from Entry
    :return: True if p is a digit, False if not
    """
    if p.isdigit():
        return True
    return False


def flag(i, com, xi, yi):
    """
    Function register click on a checkbutton and performs command
    :param i: flag number which corresponds to command number in list_cb
    :param com: command name
    :param xi: x coordinate of chk
    :param yi: y coordinate of chk
    """
    chk = Checkbutton(window1, variable=list_cb[i], command=com)
    chk.focus()
    chk.place(x=xi, y=yi)
    return


def combin():
    """
    Function reads number from question_amount and launches the desired function
    :return:
    """
    question_amount.place(x=600, y=155)
    question_amount_2.place(x=600, y=185)
    question_amount_3.place(x=600, y=215)
    question_amount_4.place(x=600, y=245)
    question_amount_5.place(x=600, y=275)
    question_amount_6.place(x=600, y=305)
    flag(1, cc_int_cc, 800, 155)
    flag(2, cc_10_in_2, 800, 185)
    flag(3, cc_2_in_10, 800, 215)
    flag(4, dop_code, 800, 245)
    flag(5, binary_arithmetic, 800, 275)
    flag(6, real_10_in_2, 800, 305)

def put_filename():
    """
    Function create 5 task and 5 answer files
    Filename enter by users
    :return:
    """
    global name1_1, name1_2, name1_3, name1_4, name1_5
    global name2_1, name2_2, name2_3, name2_4, name2_5
    nm = fn.get()
    name1_1 = nm + "1.docx"
    name1_2 = nm + "2.docx"
    name1_3 = nm + "3.docx"
    name1_4 = nm + "4.docx"
    name1_5 = nm + "5.docx"

    name2_1 = nm + "1_O.docx"
    name2_2 = nm + "2_O.docx"
    name2_3 = nm + "3_O.docx"
    name2_4 = nm + "4_O.docx"
    name2_5 = nm + "5_O.docx"

    fn.destroy()

    lbl9.destroy()
    lbl20 = Label(window1, font=("Arial Bold", 14), text='file for tasks - ' + name1_1 + '; for answers - ' + name2_1)
    lbl20.place(x=400, y=10)
    btnputinfile.destroy()
    combin()


def num_check(xi, yi, li):
    """
    Function put on a screen text if user try to put anything besides integer number
    :param xi: x coordinate for text
    :param yi: y coordinate for text
    :param li: number of command in list_cb
    """
    lb_f = Label(window1, font=("Arial Bold", 14), text='only integer number available')
    lb_f.place(x=xi, y=yi)
    list_cb[li].set(0)
    return


def dec_to_base(n, base):
    """
    Function transport n in 10 base to number in base
    :param n: number in different base
    :param base: desired base of n
    :return: number in desired base
    """
    if not hasattr(dec_to_base, 'table'):
        dec_to_base.table = '0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    n = n
    r = []
    while n:
        x, y = divmod(n, base)
        r.append(dec_to_base.table[y])
        n = x
    return ''.join(reversed(r))


def cc_int_cc():
    """
    Function put task of transfer number from first base to second in file
    """
    i2 = question_amount.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 150, 1)
    else:
        for it in range(5):
            variat(it)
            name = 'Перевод числа из одной системы счисления в другую:'
            input_file_docx(name, name)
            for p in range(int(question_amount.get())):
                number = random.randint(1, 30)
                first_base = random.randint(2, 15)
                number_in_1_base = dec_to_base(number, first_base)
                second_base = random.randint(2, 15)
                # при совпадении СС
                if first_base == second_base:
                    second_base += 1
                number_in_2_base = dec_to_base(number, second_base)
                task = f'{p + 1}. Перевести число {number_in_1_base}, записанное в {first_base} системе счисления в {second_base} систему счисления'
                answer = task + f' \n Ответ: {number_in_2_base}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=150)
    return


def cc_10_in_2():
    """
    Task transport integer numbers from 10 base to 2 base
    :return:
    """
    i2 = question_amount_2.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 180, 2)
    else:
        for it in range(5):
            variat(it)
            name = 'Перевод целых 10-ичных чисел в 2-ичную СС:'
            input_file_docx(name, name)
            for p in range(int(question_amount_2.get())):
                number = random.randint(1, 1024)
                number_in_2_base = dec_to_base(number, 2)
                task = f'{p + 1}. Перевести число {number}, записанное в 10-ичной системе счисления в 2-ичную'
                answer = task + f' \n Ответ: {number_in_2_base}'
                input_file_docx(task, answer)
            lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
            lbvi.place(x=800, y=180)
    return


def cc_2_in_10():
    """
    ask transport integer numbers from 2 base to 10 base
    :return:
    """
    i2 = question_amount_3.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 210, 3)
    else:
        for it in range(5):
            variat(it)
            name = 'Перевод целых 2-ичных чисел в 10-ичную СС:'
            input_file_docx(name, name)
            for p in range(int(question_amount_3.get())):
                number = random.randint(1, 1024)
                number_in_2_base = dec_to_base(number, 2)
                task = f'{p + 1}. Перевести число {number_in_2_base}, записанное в 2-ичной системе счисления в 10-ичную'
                answer = task + f' \n Ответ:{number}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=210)
    return


def dop(n, bits=8):
    """
    Function transfer n to addition code with 0b prefix
    :param n: Number which need to be transfered
    :param bits:
    :return:
    """
    mask = (1 << bits) - 1
    if n < 0:
        n = ((abs(n) ^ mask) + 1)
    return bin(n & mask)


def delete_0b(n):
    """
    Function remove 0b prefix from binary to dop part
    :param n:
    :return:
    """
    s1 = n.replace("0b", "")
    return s1


def dop_code(): #put numbers to additional code
    """
    Task to transfer number from right to additional code
    :return:
    """
    i2 = question_amount_4.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 240, 4)
    else:
        for it in range(5):
            variat(it)
            name = 'Представление чисел в дополнительном коде:'
            input_file_docx(name, name)
            for p in range(int(question_amount_4.get())):
                number = random.randint(-30, 10)
                task = f'{p + 1}. Представить число {number} в дополнительном коде'
                answer = task + f'\n Ответ: {delete_0b(dop(number))}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=240)
    return


def b_remove(n):
    """
    Function remove "0b" prefix
    :param n: the number to remove prefix
    :return: the number without prefix
    """
    s = bin(n)
    s1 = s.replace("0b", "")
    return s1


def binary_arithmetic():
    """
    Function put task of arithmetic operations in 2 base in file
    """
    i2 = question_amount_5.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 270, 5)
    else:
        for it in range(5):
            variat(it)
            name = 'Операции в 2-ичной СС:'
            input_file_docx(name, name)
            for p in range(int(question_amount_5.get())):
                number1 = random.randint(-10, 10)
                number2 = random.randint(-10, 10)
                sgn = random.choice('+-*')
                if number2 < 0:
                    k = 1
                else:
                    k = 0
                str_sgn = str(number1) + sgn + str(number2)
                rsh = eval(str_sgn)
                number1_without_b = b_remove(number1)
                number2_without_b = b_remove(number2)
                if k == 1:
                    number2_without_b_p = '(' + number2_without_b + ')'
                else:
                    number2_without_b_p = number2_without_b
                str_sgn_print = number1_without_b + ' ' + sgn + ' ' + number2_without_b_p
                rshb = b_remove(rsh)
                task = f'{p + 1}. Выполните следующую операцию {str_sgn_print}, ответ запишите в двоичной СС'
                answer = task + f' \n Ответ: {rshb}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=270)
    return


def dvzn(stri):
    """
    Function cuts off number up to the third decimal place
    :param stri: the float number
    :return: the float number with only three decimal place
    """
    istr = stri.find('.')
    stri1 = stri[0:(istr+4)]
    return stri1


def float_to_binary(num):
    """

    :param num:
    :return:
    """
    exponent = 0
    shifted_num = num
    while shifted_num != int(shifted_num):
        shifted_num *= 2
        exponent += 1
    if exponent == 0:
        return '{0:0b}'.format(int(shifted_num))
    binary = '{0:0{1}b}'.format(int(shifted_num), exponent+1)
    integer_part = binary[:-exponent]
    fractional_part = binary[-exponent:].rstrip('0')
    return '{0}.{1}'.format(integer_part, fractional_part)


def real_10_in_2():
    """
    Function put task  of transforming float real number from 10 base to 2
    """
    i2 = question_amount_6.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 300, 6)
    else:
        for it in range(5):
            variat(it)
            name = 'Перевод действительных 10 чисел в 2-ичную СС:'
            input_file_docx(name, name)
            for p in range(int(question_amount_6.get())):
                chis3 = int(random.uniform(-10, 10) * 100) / 100
                chis3b = float_to_binary(chis3)
                task = f'{p + 1}. Переведите число из 10-ичной СС {chis3} в 2-ичную с точность 3 знака после запятой'
                answer = task + f' \n Oтвет: {dvzn(chis3b)}'
                input_file_docx(task, answer)
            lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
            lbvi.place(x=800, y=300)
    return


window1 = Tk()


global strn
strn = ["1", "2", "3", "4", "5"]
for n in strn:
    exec('name1_' + n + "=" + "[]")
    exec('name2_' + n + "=" + "[]")
    exec('dt_' + n + '=' + "Document()")
    exec('dt1_' + n + '=' + "Document()")
    dtn = "dt_" + n
    exec("fio_add(" + dtn+")")


window1.title("Раздел 1")
window1.geometry('1200x600')
fn = Entry(window1, width=40)
fn.focus_set()
fn.place(x=390, y=30)
btnputinfile = Button(window1, text="OK", font=14, command=put_filename, foreground="#000080")
btnputinfile.place(x=750, y=20)
btnClose = Button(window1, text="Закрыть", font=14, command=window1.destroy, foreground="#000080")
btnClose.place(x=20, y=500)
lbl2 = Label(window1, font=("Arial Bold", 14), text="ТЕМА ")
lbl2.place(x=10, y=100)
lbl3 = Label(window1, font=("Arial Bold", 14), text="КОЛ-ВО ЗАДАНИЙ    ВЫБРАТЬ")
lbl3.place(x=600, y=100)

lbl1 = Label(window1, font=("Arial Bold", 14),
             text="Перевод целых чисел из одной системы счисления в другую")
lbl4 = Label(window1, font=("Arial Bold", 14),
             text="Перевод целого числа из 10-ичной в 2-ичную СС")
lbl5 = Label(window1, font=("Arial Bold", 14),
             text="Перевод целого числа из 2-ичной в 10-ичную СС")
lbl6 = Label(window1, font=("Arial Bold", 14),
             text="Перевод целого числа в дополнительный код")
lbl7 = Label(window1, font=("Arial Bold", 14),
             text="Двоичная арифметика с целыми числами")
lbl11 = Label(window1, font=("Arial Bold", 14),
              text="Перевод чисел с плавающей запятой")
lbl1.place(x=10, y=150)
lbl4.place(x=10, y=180)
lbl5.place(x=10, y=210)
lbl6.place(x=10, y=240)
lbl7.place(x=10, y=270)
lbl11.place(x=10, y=300)

question_amount = Entry(window1, width=10)
question_amount.focus_set()
question_amount_2 = Entry(window1, width=10)
question_amount_2.focus_set()
question_amount_3 = Entry(window1, width=10)
question_amount_3.focus_set()
question_amount_4 = Entry(window1, width=10)
question_amount_4.focus_set()
question_amount_5 = Entry(window1, width=10)
question_amount_5.focus_set()
question_amount_6 = Entry(window1, width=10)
question_amount_6.focus_set()

list_cb = []
for j in range(7):
    list_cb.append(IntVar())

lbl8 = Label(window1, font=("Arial Bold", 14), text="Название файла для записи")
lbl8.place(x=100, y=30)
lbl9 = Label(window1, font=("Arial Bold", 14), text=".docx")
lbl9.place(x=640, y=30)

window1.mainloop()
strn = ["1", "2", "3", "4", "5"]
for n in strn:
    exec('dt_' + n + '.save(name1_' + n + ')')
    exec('dt1_' + n + '.save(name2_' + n + ')')

exit()
