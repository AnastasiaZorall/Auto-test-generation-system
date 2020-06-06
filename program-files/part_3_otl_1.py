from docx import Document
from docx.shared import Pt
import tkinter as tk
from tkinter import *
from tkinter.ttk import Checkbutton
import random
from itertools import *
from PIL import ImageTk, Image


def combin():
    """
    Function reads number from question_amount and launches the desired function
    :return:
    """
    question_amount_1.place(x=550, y=215)
    question_amount_2.place(x=550, y=275)
    question_amount_3.place(x=550, y=335)
    question_amount_4.place(x=550, y=395)
    flag(1, comb_n_k, 800, 215)
    flag(2, comp_n_k_rep, 800, 275)
    flag(3, accom_n_k_rep, 800, 335)
    flag(4, accom_n_k, 800, 395)


def variat(it):
    """
    Function allow to work with file of variant, which you need right now
    :param it: The number of varial
    :return:
    """
    global dti, dti1
    if it == 0:
        dti = dt_1
        dti1 = dt1_1
    if it == 1:
        dti = dt_2
        dti1 = dt1_2
    if it == 2:
        dti = dt_3
        dti1 = dt1_3
    if it == 3:
        dti = dt_4
        dti1 = dt1_4
    if it == 4:
        dti = dt_5
        dti1 = dt1_5

def fio_add(dti):
    """
    Function add first/second name and patronymic in file
    No input and output data
    """
    paragraph = dti.add_paragraph('Фамилия, Имя, Группа' + '_' * 72)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(3.0)
    paragraph = dti.add_paragraph('Дата выполнения' + '_' * 80 + '\n')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2.0)


def input_file_docx(str_write, str_answer):
    """
    Function add a new paragraph in two files (main+answers)
    :param str_write: string which will be add to the main file
    :param str_answer: string which will be add to file with answers
    """
    paragraph = dti.add_paragraph(str_write)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(1.0)

    paragraph = dti1.add_paragraph(str_answer)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(1.0)


def only_int(p):
    """
    Checking string p if it is a number
    :param p: string from Entry
    :return: True if p is a digit, False if not
    """
    if p.isdigit():
        return True
    return False

def num_check(xi, yi, li):
    """
    Function put on a screen text if user try to put anything besides integer number
    :param xi: x coordinate for text
    :param yi: y coordinate for text
    :param li: number of command in list_cb
    """
    lb_f = Label(window1, font=("Arial Bold", 14), text='only integer number available')
    lb_f.place(x=xi, y=yi)
    list_cb[li].set(0)
    return


def flag(i, com, xi, yi):
    """
    Function register click on a checkbutton and performs command
    :param i: flag number which corresponds to command number in list_cb
    :param com: command name
    :param xi: x coordinate of chk
    :param yi: y coordinate of chk
    """
    chk = Checkbutton(window1, variable=list_cb[i], command=com)
    chk.focus()
    chk.place(x=xi, y=yi)
    return


def put_filename():
    """
    Function create 5 task and 5 answer files
    Filename enter by users
    :return:
    """
    global name1_1, name1_2, name1_3, name1_4, name1_5
    global name2_1, name2_2, name2_3, name2_4, name2_5
    nm = fn.get()
    name1_1 = nm + "1.docx"
    name1_2 = nm + "2.docx"
    name1_3 = nm + "3.docx"
    name1_4 = nm + "4.docx"
    name1_5 = nm + "5.docx"

    name2_1 = nm + "1_O.docx"
    name2_2 = nm + "2_O.docx"
    name2_3 = nm + "3_O.docx"
    name2_4 = nm + "4_O.docx"
    name2_5 = nm + "5_O.docx"

    fn.destroy()

    lbl9.destroy()
    lbl20 = Label(window1, font=("Arial Bold", 14), text='file for tasks - ' + name1_1 + '; for answers - ' + name2_1)
    lbl20.place(x=400, y=10)
    btnputinfile.destroy()
    combin()


def Pol_str_iz(str1):
    global strP, ik
    ik = random.randint(4, 7)
    i = 1
    strP = ''
    strr = random.choice(str1)
    strP = strP + strr
    while i <= ik - 1:
        j = 0
        strr = random.choice(str1)
        while j < len(strP):
            if strP[j] != strr:
                j += 1
            else:
                strr = random.choice(str1)
        strP = strP + strr
        i += 1


def comb_n_k():
    i2 = question_amount_1.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 215, 1)
    else:
        for it in range(5):
            variat(it)
            name = 'Формулы комбинаторики 1:'
            input_file_docx(name, name)
            str1 = "abcdefghijklmnopqrstuvwxyz"
            for p in range(int(question_amount_1.get())):
                Pol_str_iz(str1)
                dl = random.randint(3, ik - 1)
                task = f"{p + 1}. Сколько существует способов выбрать {dl} букв из строки {strP} (повторяющиеся буквы считаются разными) ?"
                a = len(list(combinations(strP, dl)))
                answer = task + '\n' + f'Ответ:\n {a}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено)                                  ")
        lbvi.place(x=800, y=215)

def accom_n_k():
    i2 = question_amount_4.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 395, 4)
    else:
        for it in range(5):
            variat(it)
            name = 'Формулы комбинаторики 3:'
            input_file_docx(name, name)

            str1 = "abcdefghijklmnopqrstuvwxyz"
            for p in range(int(question_amount_4.get())):
                Pol_str_iz(str1)
                dl = random.randint(3, ik - 1)
                task = f"{p + 1}. Сколько {dl}-значных шифров можно {strP} из строки при условии, что знаки в шифре могут повторяться"
                k1 = 0
                r1 = product(strP, repeat=dl)
                for var in r1:
                    k1 += 1
                answer = task + '\n' + f'Ответ:\n {k1}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=395)

def accom_n_k_rep():
    i2 = question_amount_3.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 335, 3)
    else:
        for it in range(5):
            variat(it)
            name = 'Формулы комбинаторики 4:'
            input_file_docx(name, name)
            str1 = "abcdefghijklmnopqrstuvwxyz"
            for p in range(int(question_amount_3.get())):
                Pol_str_iz(str1)
                dl = random.randint(3, ik - 1)
                task = f"{p + 1}. Сколько {dl}-значных шифров можно {strP} из строки при условии, что знаки в шифре не повторяются"
                k2 = 0
                for var in permutations(strP, dl):
                    k2 += 1
                answer = task + '\n' + f'Ответ:\n {k2}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=335)

def comp_n_k_rep():
    i2 = question_amount_2.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 275, 2)
    else:
        for it in range(5):
            variat(it)
            name = 'Формулы комбинаторики 2:'
            input_file_docx(name, name)
            str1 = "abcdefghijklmnopqrstuvwxyz"
            for p in range(int(question_amount_2.get())):
                Pol_str_iz(str1)
                dl = random.randint(2, ik - 1)
                a = len(list(combinations_with_replacement(strP, dl)))
                task = f"{p + 1}. Сколько существует способов выбрать {dl} букв из строки {strP} необходимо учесть повторы?"
                answer = task + '\n' + f'Ответ:\n {a}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=275)

window1 = Tk()
global strn
strn = ["1","2","3","4","5"]
for n in strn:
    exec ('name1_'+ n + "=" +"[]")
    exec ('name2_'+ n + "=" +"[]")
    exec ('dt_' + n+ '=' + "Document()")
    exec ('dt1_' + n+ '=' + "Document()")
    dtn = "dt_" +n
    exec("fio_add(" + dtn+")")

window1.title("Раздел 3")
window1.geometry('1200x600')
fn = Entry(window1, width=40)
fn.focus_set()
fn.place(x=390, y=30)

btnputinfile = Button(window1, text="OK", font=14, command=put_filename, foreground="#000080")
btnputinfile.place(x=750, y=20)
btnClose = Button(window1, text="Закрыть", font=14, command=window1.destroy, foreground="#000080")
btnClose.place(x=20, y=500)
lbl2 = Label(window1, font=("Arial Bold", 14), text="ТЕМА ")
lbl2.place(x=10, y=100)
lbl3 = Label(window1, font=("Arial Bold", 14), text="КОЛ-ВО ЗАДАНИЙ    ВЫБРАТЬ")
lbl3.place(x=550, y=100)

lbl4 = Label(window1, font=("Arial Bold", 14), text="Сочетания без повторений")
lbl5 = Label(window1, font=("Arial Bold", 14), text="Сочетания c повторениями")
lbl6 = Label(window1, font=("Arial Bold", 14), text="Размещения с повторениями")
lbl7 = Label(window1, font=("Arial Bold", 14), text="Размещения без повторений")
lbl4.place(x=10, y=210)
lbl5.place(x=10, y=270)
lbl6.place(x=10, y=330)
lbl7.place(x=10, y=390)

question_amount_1 = Entry(window1, width=10)
question_amount_1.focus_set()

question_amount_2 = Entry(window1, width=10)
question_amount_2.focus_set()

question_amount_3 = Entry(window1, width=10)
question_amount_3.focus_set()

question_amount_4 = Entry(window1, width=10)
question_amount_4.focus_set()
list_cb = []
for j in range(5):
    list_cb.append(IntVar())
lbl8 = Label(window1, font=("Arial Bold", 14), text="Название файла для записи")
lbl8.place(x=100, y=30)
lbl9 = Label(window1, font=("Arial Bold", 14), text=".docx")
lbl9.place(x=640, y=30)

#path1 = 'C:/University/2_course_work/cnm.jpg'
path1 = 'cnm.jpg'

img1 = ImageTk.PhotoImage(Image.open(path1))
panel = tk.Label(window1, image=img1)
panel.pack(side="bottom", fill="both", expand="yes")
panel.place(x=300, y=200)
path2 = 'cnmc.jpg'

#path2 = 'C:/University/2_course_work/cnmc.jpg'
img2 = ImageTk.PhotoImage(Image.open(path2))
panel2 = tk.Label(window1, image=img2)
panel2.pack(side="bottom", fill="both", expand="yes")
panel2.place(x=300, y=260)
path3 = 'rcp.jpg'

#path3 = 'C:/University/2_course_work/rcp.jpg'
img3 = ImageTk.PhotoImage(Image.open(path3))
panel3 = tk.Label(window1, image=img3)
panel3.pack(side="bottom", fill="both", expand="yes")
panel3.place(x=300, y=320)
path4 = 'rbp.jpg'

#path4 = 'C:/University/2_course_work/rbp.jpg'
img4 = ImageTk.PhotoImage(Image.open(path4))
panel4 = tk.Label(window1, image=img4)
panel4.pack(side="bottom", fill="both", expand="yes")
panel4.place(x=300, y=380)

window1.mainloop()
strn = ["1","2","3","4","5"]
for n in strn:
    exec ('dt_' + n+ '.save(name1_' + n + ')')
    exec ('dt1_' + n+ '.save(name2_' + n + ')')

exit()
