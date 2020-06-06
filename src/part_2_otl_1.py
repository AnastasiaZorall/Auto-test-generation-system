import docx
from docx import Document
from docx.shared import Pt
from tkinter import *
from tkinter.ttk import Checkbutton
import random
import numpy as np
import pandas as pd


def variat(it):
    """
    Function allow to work with file of variant, which you need right now
    :param it: The number of varial
    :return:
    """
    global dti, dti1
    if it == 0:
        dti = dt_1
        dti1 = dt1_1
    if it == 1:
        dti = dt_2
        dti1 = dt1_2
    if it == 2:
        dti = dt_3
        dti1 = dt1_3
    if it == 3:
        dti = dt_4
        dti1 = dt1_4
    if it == 4:
        dti = dt_5
        dti1 = dt1_5


def fio_add(dti):
    """
    Function add first/second name and patronymic in file
    No input and output data
    """
    paragraph = dti.add_paragraph('Фамилия, Имя, Группа' + '_' * 72)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(3.0)
    paragraph = dti.add_paragraph('Дата выполнения' + '_' * 80 + '\n')
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(2.0)


def input_file_docx(str_write, str_answer):
    """
    Function add a new paragraph in two files (main+answers)
    :param str_write: string which will be add to the main file
    :param str_answer: string which will be add to file with answers
    """
    paragraph = dti.add_paragraph(str_write)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(1.0)

    paragraph = dti1.add_paragraph(str_answer)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(1.0)


def only_int(p):
    """
    Checking string p if it is a number
    :param p: string from Entry
    :return: True if p is a digit, False if not
    """
    if p.isdigit():
        return True
    return False


def num_check(xi, yi, li):
    """
    Function put on a screen text if user try to put anything besides integer number
    :param xi: x coordinate for text
    :param yi: y coordinate for text
    :param li: number of command in list_cb
    """
    lb_f = Label(window1, font=("Arial Bold", 14), text='only integer number available')
    lb_f.place(x=xi, y=yi)
    list_cb[li].set(0)
    return


def flag(i, com, xi, yi):
    """
    Function register click on a checkbutton and performs command
    :param i: flag number which corresponds to command number in list_cb
    :param com: command name
    :param xi: x coordinate of chk
    :param yi: y coordinate of chk
    """
    chk = Checkbutton(window1, variable=list_cb[i], command=com)
    chk.focus()
    chk.place(x=xi, y=yi)
    return


def combin():
    """
    Function reads number from question_amount and launches the desired function
    :return:
    """
    question_amount_1.place(x=550, y=215)
    question_amount_2.place(x=550, y=245)
    question_amount_3.place(x=550, y=275)
    question_amount_4.place(x=550, y=305)
    question_amount_5.place(x=550, y=335)

    flag(1, boolean, 800, 215)
    flag(2, cdnf, 800, 245)
    flag(3, cknf, 800, 275)
    flag(4, mdnf, 800, 305)
    flag(5, mknf, 800, 335)


def put_filename():
    """
    Function create 5 task and 5 answer files
    Filename enter by users
    :return:
    """
    global name1_1, name1_2, name1_3, name1_4, name1_5
    global name2_1, name2_2, name2_3, name2_4, name2_5
    nm = fn.get()
    name1_1 = nm + "1.docx"
    name1_2 = nm + "2.docx"
    name1_3 = nm + "3.docx"
    name1_4 = nm + "4.docx"
    name1_5 = nm + "5.docx"

    name2_1 = nm + "1_O.docx"
    name2_2 = nm + "2_O.docx"
    name2_3 = nm + "3_O.docx"
    name2_4 = nm + "4_O.docx"
    name2_5 = nm + "5_O.docx"

    fn.destroy()

    lbl9.destroy()
    lbl20 = Label(window1, font=("Arial Bold", 14), text='file for tasks - ' + name1_1 + '; for answers - ' + name2_1)
    lbl20.place(x=400, y=10)
    btnputinfile.destroy()
    combin()


def boolean():
    """
    Boolean algebra task. Included addition subtraction multiplication
    :return:
    """
    i2 = question_amount_1.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 210, 0)
    else:
        for it in range(5):
            variat(it)
            name = 'Операции булевой алгебры:'
            input_file_docx(name, name)
            for p in range(int(question_amount_1.get())):
                spechotr = ''
                sgnA = random.randint(0, 1)
                print()
                sgnB = random.randint(0, 1)
                print ("A =", sgnA, "B =", sgnB)
                spech = ''
                strv = ''
                ir = 1
                while ir < 4:
                    nch = random.randint(1, 4)
                    if nch == 1:
                        spech = spech + "¬A∨B"
                        strv = strv + "not sgnA or sgnB "
                    if nch == 2:
                        strv = strv + "not sgnB or sgnA"
                        spech = spech + "¬B∨A"
                    if nch == 3:
                        strv = strv + "(sgnB or sgnA)"
                        spech = spech + "(B∨A)"
                    if nch == 4:
                        strv = strv + "(sgnA and sgnB)"
                        spech = spech + "(A∧B)"
                    if ir < 3:
                        strv = strv + " and "
                        spech = spech + "∧"
                    ir = ir + 1
                task = f'{p + 1}. A = {sgnA}, B = {sgnB} Какое значение принимает функция {spech} ?'
                rsh = eval(strv)
                answer = task + f' \n Oтвет: {rsh}'
                input_file_docx(task, answer)
            lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
            lbvi.place(x=800, y=210)
    return


def cknf():
    """
    Task to create perfect conjunctive function from function table
    :return:
    """
    i2 = question_amount_3.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 270, 3)
    else:
        for it in range(5):
            variat(it)
            name = 'СKНФ:'
            input_file_docx(name, name)
            for p in range(int(question_amount_3.get())):
                df = pd.DataFrame(np.array([[0, 0, 0],
                                            [0, 0, 1],
                                            [0, 1, 0],
                                            [0, 1, 1],
                                            [1, 0, 0],
                                            [1, 0, 1],
                                            [1, 1, 0],
                                            [1, 1, 1]]), columns=['x', 'y', 'z'])
                df["F(x, y, z)"] = np.random.randint(0, 2, size=len(df))
                task = f'Постойте СКНФ для следующей таблицы: \n {df}'
                r = (df.loc[df['F(x, y, z)'] == 0, ['x', 'y', 'z']].astype(bool)).astype('int8')
                res = (r.apply(lambda r: '({}{} v {}{} v {}{})'.format('!' * r['x'], 'x', '!' * r['y'],
                                                                       'y', '!' * r['z'], 'z'), axis=1).str.cat(sep=' ^ '))
                answer = task + f'\n Ответ: {res}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=270)

def cdnf():
    """
    Task to create perfect disjunctive function from function table
    :return:
    """
    i2 = question_amount_2.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 240, 2)
    else:
        for it in range(5):
            variat(it)
            name = 'СДНФ:'
            input_file_docx(name, name)
            for p in range(int(question_amount_2.get())):
                df = pd.DataFrame(np.array([[0, 0, 0],
                                            [0, 0, 1],
                                            [0, 1, 0],
                                            [0, 1, 1],
                                            [1, 0, 0],
                                            [1, 0, 1],
                                            [1, 1, 0],
                                            [1, 1, 1]]), columns=['x', 'y', 'z'])
                df["F(x, y, z)"] = np.random.randint(0, 2, size=len(df))
                task = f'Постойте СДНФ для следующей таблицы: \n {df}'
                r = (~df.loc[df['F(x, y, z)'] == 1, ['x', 'y', 'z']].astype(bool)).astype('int8')
                res = (r.apply(lambda r: '({}{} ^ {}{} ^ {}{})'.format('!' * r['x'], 'x', '!' * r['y'],
                                                                       'y', '!' * r['z'], 'z'), axis=1).str.cat(sep=' v '))
                answer = task + f'\n Ответ: {res}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=240)


def mdnf():
    """
    Task to create minimal disjunctive function from function table
    :return:
    """
    i2 = question_amount_4.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 300, 4)
    else:
        for it in range(5):
            variat(it)
            name = 'МДНФ:'
            input_file_docx(name, name)
            for p in range(int(question_amount_4.get())):
                df = pd.DataFrame(np.array([[0, 0, 0],
                                            [0, 0, 1],
                                            [0, 1, 0],
                                            [0, 1, 1],
                                            [1, 0, 0],
                                            [1, 0, 1],
                                            [1, 1, 0],
                                            [1, 1, 1]]), columns=['x', 'y', 'z'])
                iNomer = random.randint(1, 10)
                doc_O = docx.Document('mdnf.docx')
                doc_t = docx.Document('f_meaning.docx')
                strW = doc_t.paragraphs[iNomer - 1].text
                a = strW.split(',')
                df["F(x, y, z)"] = a
                task = f'Построить МДНФ по таблице истинности \n{df}'
                strW_Otvet = doc_O.paragraphs[iNomer - 1].text
                answer = task + f' \n Oтвет: \n {strW_Otvet}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=300)


def mknf():
    """
    Task to create perfect conjunctive function from function table
    :return:
    """
    i2 = question_amount_5.get()
    check = (only_int(i2))
    if not check:
        num_check(820, 330, 5)
    else:
        for it in range(5):
            variat(it)
            name = 'МКНФ:'
            input_file_docx(name, name)
            for p in range(int(question_amount_5.get())):
                df = pd.DataFrame(np.array([[0, 0, 0],
                                            [0, 0, 1],
                                            [0, 1, 0],
                                            [0, 1, 1],
                                            [1, 0, 0],
                                            [1, 0, 1],
                                            [1, 1, 0],
                                            [1, 1, 1]]), columns=['x', 'y', 'z'])
                iNomer = random.randint(1, 10)
                doc_O = docx.Document('mknf.docx')
                doc_t = docx.Document('f_meaning.docx')
                strW = doc_t.paragraphs[iNomer - 1].text
                a = strW.split(',')
                df["F(x, y, z)"] = a
                task = f'Построить МКНФ по таблице истинности \n{df}'
                strW_Otvet = doc_O.paragraphs[iNomer - 1].text
                answer = task + f' \n Oтвет: \n {strW_Otvet}'
                input_file_docx(task, answer)
        lbvi = Label(window1, font=("Arial Bold", 14), text="Выполнено                                  ")
        lbvi.place(x=800, y=330)


window1 = Tk()

global strn
strn = ["1", "2", "3", "4", "5"]
for n in strn:
    exec ('name1_' + n + "=" +"[]")
    exec ('name2_' + n + "=" +"[]")
    exec ('dt_' + n + '=' + "Document()")
    exec ('dt1_' + n + '=' + "Document()")
    dtn = "dt_" + n
    exec("fio_add(" + dtn+")")

window1.title("Раздел 1")
window1.geometry('1200x600')
fn = Entry(window1, width=40)
fn.focus_set()
fn.place(x=390, y=30)
btnputinfile = Button(window1, text="OK", font=14, command=put_filename, foreground="#000080")
btnputinfile.place(x=750, y=20)
btnClose = Button(window1, text="Закрыть", font=14, command=window1.destroy, foreground="#000080")
btnClose.place(x=20, y=500)
lbl2 = Label(window1, font=("Arial Bold", 14), text="ТЕМА ")
lbl2.place(x=10, y=100)
lbl3 = Label(window1, font=("Arial Bold", 14), text="КОЛ-ВО ЗАДАНИЙ    ВЫБРАТЬ")
lbl3.place(x=550, y=100)

lbl4 = Label(window1, font=("Arial Bold", 14), text="Операции булевой алгебры")
lbl6 = Label(window1, font=("Arial Bold", 14), text="СДНФ")
lbl7 = Label(window1, font=("Arial Bold", 14), text="СКНФ")
lbl10 = Label(window1, font=("Arial Bold", 14), text="МДНФ")
lbl11 = Label(window1, font=("Arial Bold", 14), text="МКНФ")
lbl4.place(x=10, y=210)
lbl6.place(x=10, y=240)
lbl7.place(x=10, y=270)
lbl10.place(x=10, y=300)
lbl11.place(x=10, y=330)


question_amount_1 = Entry(window1, width=10)
question_amount_1.focus_set()

question_amount_2 = Entry(window1, width=10)
question_amount_2.focus_set()

question_amount_3 = Entry(window1, width=10)
question_amount_3.focus_set()

question_amount_4 = Entry(window1, width=10)
question_amount_4.focus_set()

question_amount_5 = Entry(window1, width=10)
question_amount_5.focus_set()

list_cb = []
for j in range(7):
    list_cb.append(IntVar())

lbl8 = Label(window1, font=("Arial Bold", 14), text="Название файла для записи")
lbl8.place(x=100, y=30)
lbl9 = Label(window1, font=("Arial Bold", 14), text=".docx")
lbl9.place(x=640, y=30)

window1.mainloop()
strn = ["1", "2", "3", "4", "5"]
for n in strn:
    exec('dt_' + n + '.save(name1_' + n + ')')
    exec('dt1_' + n + '.save(name2_' + n + ')')
exit()
